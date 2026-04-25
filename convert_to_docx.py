#!/usr/bin/env python
"""Convert DESCRIPTION.txt to description.docx"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    # Read text file
    with open('DESCRIPTION.txt', 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    # Create DOCX
    doc = Document()
    
    # Parse content and add to doc
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            doc.add_paragraph()
        elif stripped.startswith('═') or stripped.startswith('━'):
            continue
        elif stripped.isupper() and len(stripped) > 20 and ' - ' in stripped:
            # Section title
            heading = doc.add_heading(stripped, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif stripped.startswith('✓') or stripped.startswith('•') or stripped.startswith('-'):
            # Bullet point
            doc.add_paragraph(stripped, style='List Bullet')
        elif stripped[0].isdigit() and (stripped.startswith('1. ') or '4. ' in stripped[:5]):
            # Numbered list
            doc.add_paragraph(stripped, style='List Number')
        else:
            # Regular paragraph
            doc.add_paragraph(stripped, style='Normal')
    
    # Save
    doc.save('description.docx')
    print("✅ SUCCESS: description.docx created!")
    print("📁 File: c:\\Users\\Rizwan Shaikh\\Desktop\\RAGnosis\\description.docx")
    
except Exception as e:
    print(f"❌ Error: {e}")
    import traceback
    traceback.print_exc()
